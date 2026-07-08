"""生成"AI分类算法'黑箱'可视化"操作手册 docx"""

import json
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 读取当前配置 ──
CONFIG_PATH = os.path.join(os.path.dirname(__file__), "config.json")
with open(CONFIG_PATH, "r", encoding="utf-8") as f:
    CFG = json.load(f)

SCENE = CFG["scene_name"]
TITLE = CFG["title"]
CAPTION = CFG["caption"]
POSITIVE = CFG["positive_label"]
NEGATIVE = CFG["negative_label"]
FEATURE_X = CFG["feature_x"]
FEATURE_Y = CFG["feature_y"]

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
                           "操作手册：AI分类算法黑箱可视化（LLM增强版）.docx")


# ── 辅助函数 ──
def set_spacing(para, before=0, after=0, line=None):
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    if before:
        spacing.set(qn('w:before'), str(before))
    if after:
        spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))


def add_run(para, text, font_name='宋体', size=Pt(12),
            bold=False, color=RGBColor(0, 0, 0)):
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    return run


def add_heading_para(text, level=1, center=False):
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(para, before=200 if level <= 2 else 120, after=80)
    sizes = {1: Pt(22), 2: Pt(16), 3: Pt(14), 4: Pt(12)}
    add_run(para, text, font_name='黑体',
            size=sizes.get(level, Pt(14)), bold=True)
    return para


def add_body(text, bold=False, indent=False, size=Pt(12)):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Cm(0.75)
    set_spacing(para, before=0, after=0, line=360)
    add_run(para, text, bold=bold, size=size)
    return para


def add_label_value(label, value='', value_bold=False):
    para = doc.add_paragraph()
    set_spacing(para, before=0, after=0, line=360)
    add_run(para, label, bold=True)
    if value:
        add_run(para, value, bold=value_bold)
    return para


def add_bullet(text, indent_level=0, bold_prefix=''):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1 + indent_level * 0.6)
    set_spacing(para, before=0, after=0, line=340)
    if bold_prefix:
        add_run(para, bold_prefix, bold=True)
    add_run(para, text)
    return para


def add_blank_line():
    para = doc.add_paragraph()
    set_spacing(para, before=0, after=0)
    add_run(para, '')
    return para


# ── 开始构建文档 ──
doc = Document()

# 全局默认样式
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0, 0, 0)
rFonts = style.element.rPr.rFonts
rFonts.set(qn('w:eastAsia'), '宋体')

# ===== 文档标题 =====
add_heading_para('智能体操作手册', level=1, center=True)
add_blank_line()

# ===== 1. 智能体基本信息 =====
add_heading_para('一、智能体基本信息', level=2)

add_label_value('智能体名称：',
                'AI分类算法"黑箱"可视化——多模型决策边界对比探究智能体',
                value_bold=True)
add_label_value('当前应用场景：', SCENE)
add_label_value('作者：', '吴明钊')
add_label_value('作者单位：', '茂名市高新中学')
add_label_value('适用学段：', '初中、高中信息技术/人工智能课程')
add_label_value('智能体类型：', '助学智能体')
add_blank_line()

# ===== 2. 应用访问地址及账号 =====
add_heading_para('二、应用访问地址及账号', level=2)

add_label_value('访问方式一（在线版）：')
add_body('https://airobustnesslab-qpyhrdy4d6ndzea6zrniyy.streamlit.app/')
add_body('（注：在线版部署于 Streamlit Community Cloud，首次加载约需 10–30 秒唤醒服务器）')
add_blank_line()
add_label_value('访问方式二（本地运行）：')
add_body('1. 确保已安装 Python 3.8+ 环境')
add_body('2. 在项目目录下执行：pip install -r requirements.txt')
add_body('3. 执行：streamlit run app.py')
add_body('4. 浏览器自动打开 http://localhost:8503')
add_blank_line()
add_label_value('体验账号：', '无需注册登录，打开即用')
add_blank_line()

# ===== 3. 功能介绍 =====
add_heading_para('三、功能介绍', level=2)

add_body('本智能体是一款面向中小学人工智能课程的交互式可视化探究工具。'
         '它将抽象的 AI 分类算法原理转化为直观的"决策边界对比图"，'
         '让学习者通过调节参数、观察变化、记录数据的方式，'
         '自主探究四种经典分类算法（KNN、决策树、SVM、逻辑回归）'
         '在相同数据上的不同表现。')

add_blank_line()
add_body('核心功能包括：', bold=True)

add_bullet('多模型同台对比：同一份训练数据，同时训练 KNN、决策树、'
           'SVM、逻辑回归四种算法，生成 2×2 四格决策边界对比图，'
           '让学习者直观比较不同算法的"思维差异"。',
           bold_prefix='（1）')

add_bullet(f'虚实联动机制：左侧面板显示待检样本图（如荔枝照片），'
           f'学习者调节「图像模糊」「局部遮挡」滑块后，'
           f'图片发生相应降质，同时右侧四张图中代表送检样本的'
           f'"⭐"星自动发生坐标漂移，模拟 AI 特征提取的偏差。',
           bold_prefix='（2）')

add_bullet('决策边界可视化：四种算法各自的"分类决策面"以彩色等高线'
           '形式呈现在散点图上，红蓝区域分别代表 AI 的判定结果，'
           '让抽象的数学分类过程变得一目了然。',
           bold_prefix='（3）')

add_bullet('共识仪表盘：实时统计并展示四种算法的判定一致性——'
           '"全体一致""出现分歧""全体误判"三态结论，'
           '引导学生理解"多模型投票"的重要性。',
           bold_prefix='（4）')

add_bullet('探究实验台：支持一键记录当前实验参数与结果，'
           '自动生成实验数据表格，并可导出为 CSV 文件，'
           '便于学生撰写科学探究报告。',
           bold_prefix='（5）')

add_bullet('零代码场景迁移：教师只须修改 config.json 中的场景名称、'
           '特征标签、判定标签和样本图片路径，即可将整个智能体'
           '切换至任何教学场景（农作物病害、水果分级、茶叶鉴别等），'
           '无需修改任何 Python 代码。',
           bold_prefix='（6）')

add_bullet('AI 实验解读（新增）：点击「🎯 解读当前实验」按钮，'
           '智能体自动将当前实验参数（模糊度、遮挡率、重叠度、投毒率）'
           '和四种模型的判定结果发送至国产大模型（默认使用通义千问），'
           '生成一段通俗易懂的中文解读，帮助学生理解"为什么会出现这个结果"。',
           bold_prefix='（7）')

add_bullet('AI 问答助教（新增）：侧边栏常驻 AI 问答助教窗口，'
           '学生可随时输入关于算法原理、实验现象的疑问，'
           '大模型结合当前实验状态实时回答。无需注册，即问即答。',
           bold_prefix='（8）')

add_blank_line()

# ===== 4. 操作步骤 =====
add_heading_para('四、操作步骤', level=2)

# 4.1
add_heading_para('1. 启动与登录', level=3)
add_body('本智能体基于 Streamlit Web 框架构建，无需注册登录。'
         '双击项目目录中的「一键启动.bat」（Windows）或在终端执行'
         ' streamlit run app.py，浏览器自动弹出即进入主界面。')
add_blank_line()

# 4.2
add_heading_para('2. 主界面布局概览', level=3)
add_body('主界面分为四大区域：', bold=True)
add_bullet('左侧区域（现象探秘）：显示待检样本图，'
           '下方为「模型共识仪表盘」，展示四种算法的判定结果和准确率。')
add_bullet('右侧区域（机制解密）：2×2 四格决策边界对比图，'
           '可通过展开面板查看各模型详细指标和核心原理说明。')
add_bullet('侧边栏（实验环境控制台）：集中控制所有实验参数，'
           '包括图像干扰和结构化数据扰动两大类。')
add_bullet('侧边栏底部（AI 问答助教）：💬 AI 问答助教常驻显示，'
           '学生可随时提问，大模型实时解答。')
add_blank_line()

# 4.3 操作流程
add_heading_para('3. 操作流程', level=3)

add_body('步骤一：选择探究主题', bold=True)
add_body(f'当前场景为「{SCENE}」。教师可提前配置 config.json '
         f'切换至不同主题（如水稻病害检测、茶叶品质鉴别等），'
         f'学生打开页面即可开始探究。', indent=True)
add_blank_line()

add_body('步骤二：调节「非结构化感知破坏器」', bold=True)
add_body('在侧边栏第①组滑块中，向右拖动「图像模糊干扰」滑块，'
         '观察左侧图片逐渐模糊，同时右侧四张图中的 "⭐" 星'
         '开始向左下侧（危险区域）漂移。同样地，'
         '拖动「局部遮挡干扰」滑块可模拟图像被部分遮挡的效果。',
         indent=True)
add_blank_line()

add_body('步骤三：调节「结构化数据破坏器」', bold=True)
add_body('在侧边栏第②组滑块中，「特征重叠度」控制两类样本'
         '在特征空间中的混叠程度——拉高则好次品特征相近，'
         'AI 更难区分。「标注投毒率」模拟训练数据中'
         '部分标签被标注错误的情况——拉高则 AI 学到错误规律。',
         indent=True)
add_blank_line()

add_body('步骤四：观察四模型对比图', bold=True)
add_body('右侧 2×2 四格图实时展示 KNN、决策树、SVM、逻辑回归'
         '四种算法在当前数据下的决策边界。每张子图右上角标注'
         '该模型的准确率和判定结果。当 "⭐" 星跨过红蓝分界线时，'
         '表示该模型发生了误判（红色叉号标记），'
         '同时左侧样本图上触发红色印章掉落动画，'
         '产生极具视觉冲击力的反馈效果。', indent=True)
add_blank_line()

add_body('步骤五：查看共识仪表盘', bold=True)
add_body('左侧底部的「模型共识仪表盘」以 ✅/❌ 指示灯的形式'
         '直观展示四种算法各自的判定结果和准确率。'
         '下方文字总结给出"全体一致""出现分歧""全体误判"'
         '三段式结论，帮助学生快速把握当前实验的整体态势。',
         indent=True)
add_blank_line()

add_body('步骤六：记录与导出实验数据', bold=True)
add_body('点击页面底部的「📸 记录当前实验数据」按钮，'
         '可将当前参数组合和四模型的判定结果保存至实验记录表。'
         '积累多条记录后，点击「📥 导出实验报告」按钮，'
         '可下载 CSV 格式的实验数据，用 Excel 打开即可'
         '作为科学探究报告的原始数据。', indent=True)
add_blank_line()

add_body('步骤七：AI 实验解读', bold=True)
add_body('在右侧面板底部点击「🎯 解读当前实验」按钮，'
         '系统自动将当前实验参数和结果发送至国产大模型，'
         '生成一段通俗的中文解读。解读内容显示在按钮下方，'
         '帮助学生理解实验现象背后的 AI 原理。'
         '点击「🗑️ 清除解读」可清除当前解读内容。',
         indent=True)
add_blank_line()

add_body('步骤八：AI 问答助教', bold=True)
add_body('侧边栏底部「💬 AI 问答助教」区域常驻显示，'
         '学生在实验过程中如有疑问，可直接在输入框中输入问题，'
         '点击「💬 发送」按钮，大模型将结合当前实验状态'
         '给出回答。对话历史会保留在当前会话中，'
         '点击「🗑️」按钮可清空对话记录。',
         indent=True)
add_blank_line()

# 4.4
add_heading_para('4. 智能体中的插件及工作流介绍', level=3)
add_body('本智能体为混合架构，核心可视化部分为纯 Python Streamlit 应用，'
         'AI 解读与问答功能调用国产大模型 API（默认支持 Ollama 本地模型'
         '或 SiliconFlow/DeepSeek 等在线 API）。共包含三条工作流：',
         indent=True)
add_blank_line()

add_body('工作流一：核心可视化引擎（本地运行，不依赖网络）', bold=True)
add_body('用户调节侧边栏参数（模糊/遮挡/重叠/投毒）'
         '→ sklearn 生成训练数据 → 四种分类器同时训练'
         '→ 网格采样绘制决策边界 → Pillow 处理图像降质'
         '→ 测试点坐标漂移计算 → Matplotlib 输出 2×2 对比图'
         '→ 共识统计 → Streamlit 渲染界面', indent=True)
add_blank_line()

add_body('工作流二：AI 实验解读（需 LLM 后端）', bold=True)
add_body('用户点击「🎯 解读当前实验」按钮'
         '→ 系统收集当前参数（模糊/遮挡/重叠/投毒/样本数）'
         '和四模型结果（准确率/判定/是否误判）'
         '→ 组装 Prompt 发送至大模型 API（Ollama/SiliconFlow/DeepSeek）'
         '→ 大模型返回中文解读文本'
         '→ 显示在解读结果区域', indent=True)
add_blank_line()

add_body('工作流三：AI 问答助教（需 LLM 后端）', bold=True)
add_body('学生在侧边栏输入框中输入问题'
         '→ 系统将问题与当前实验参数拼接为 Prompt'
         '→ 发送至大模型 API → 大模型结合实验上下文回答'
         '→ 答案追加到对话历史 → 渲染显示', indent=True)
add_blank_line()

add_body('技术栈：', bold=True)
add_bullet('前端交互：Streamlit')
add_bullet('算法引擎：scikit-learn（KNN / 决策树 / SVM / 逻辑回归）')
add_bullet('图像处理：Pillow（高斯模糊、局部遮挡）')
add_bullet('数据可视化：Matplotlib + NumPy')
add_bullet('数据管理：Pandas（实验记录与 CSV 导出）')
add_bullet('大模型接入：httpx（兼容 Ollama / SiliconFlow / DeepSeek 等国产大模型 API）')
add_blank_line()

# 4.5
add_heading_para('5. 注意事项', level=3)

add_bullet('在线版首次加载需要 10–30 秒唤醒服务器，请耐心等待。')
add_bullet('本地运行请确保 Python 3.8+ 环境，完整安装 requirements.txt 中的依赖。')
add_bullet('侧边栏「高级参数」中的「样本数量」一般不建议修改，'
           '保持默认值（100）可获得最佳演示效果。')
add_bullet('修改 config.json 切换场景时，需同时替换 assets/ 目录下的'
           '样本图片（建议 4:3 比例），否则页面会提示图片缺失。')
add_bullet('实验记录数据仅在当前浏览器会话中有效，刷新页面前请记得导出 CSV。')
add_bullet('AI 实验解读与 AI 问答助教功能需配置大模型后端。'
           '在 config.json 的 llm 字段中设置 provider（支持 ollama'
           '本地模型 / siliconflow / deepseek 在线 API）、api_url、api_key'
           '和 model。使用 Ollama 本地模型无需 API Key，'
           '只需安装 Ollama 并拉取模型（推荐 qwen2.5:7b）。'
           '未配置 LLM 时，AI 功能按钮将提示配置信息，不影响核心可视化功能。')
add_blank_line()

# ===== 5. 价值示范 =====
add_heading_para('五、价值示范', level=2)

add_heading_para('1. 背景与适用教育情境', level=3)
add_body('在《义务教育信息科技课程标准（2022 年版）》和'
         '《普通高中信息技术课程标准（2017 年版 2020 年修订）》中，'
         '人工智能原理及其局限性被列为重要教学内容。然而，'
         'AI 分类算法的"黑箱"特性使得学生难以理解模型内部的'
         '决策机制，传统教学往往停留在"调参看准确率"的表面层面。')
add_blank_line()
add_body('本智能体适用于以下教学情境：', bold=True)
add_bullet('初中"人工智能初步"模块：理解 AI 分类的基本原理')
add_bullet('高中"人工智能基础"模块：对比不同分类算法的特性与适用场景')
add_bullet('高中"算法与程序设计"模块：理解训练数据质量对模型性能的影响')
add_bullet('信息技术/科学课外拓展活动：探究式项目学习')
add_blank_line()

add_heading_para('2. 设计意图及思路', level=3)
add_body('本作品名为「AI分类算法"黑箱"可视化——多模型决策边界对比探究智能体」，'
         '核心设计理念围绕"打开AI黑箱"展开，包含以下三个层次：', indent=True)
add_blank_line()

add_body('层次一：打开算法黑箱——"双轨空间碰撞"', bold=True, indent=True)
add_body('第一轨——"看走眼"（感性特征漂移）：通过图像模糊和局部遮挡，'
         '模拟现实世界中输入数据的降质过程，让学生直观看到'
         '"图片变糊 → 特征坐标漂移 → AI 误判"的完整因果链。',
         indent=True)
add_blank_line()
add_body('第二轨——"画错线"（理性边界扭曲）：通过数据重叠和标注投毒，'
         '模拟训练数据质量的下降，让学生亲眼见证'
         '"数据变差 → 决策边界扭曲 → AI 画错线"的渐变过程。',
         indent=True)
add_blank_line()
add_body('两条轨道在画面上交汇——当漂移的 "⭐" 星恰好跨过扭曲的'
         '决策边界时，触发极具视觉冲击力的红色印章掉落动效，'
         '深刻揭示"数据质量 + 算法局限 → 双重灾难"的核心教学主题。',
         indent=True)
add_blank_line()

add_body('层次二：多模型对比——从"一言堂"到"百家争鸣"', bold=True, indent=True)
add_body('本智能体在同一数据上同时运行四种不同算法（KNN、决策树、'
         'SVM、逻辑回归），生成 2×2 四格决策边界对比图。'
         '每种算法的"思考方式"不同——KNN 看邻居投票、决策树按规则'
         '问到底、SVM 找最宽分界线、逻辑回归算概率定阈值——'
         '面对相同的干扰，有的模型扛住了，有的误判了。'
         '通过这种"同台竞技"的设计，学生能直观理解：'
         '没有"最好"的算法，不同算法有各自的优劣势和鲁棒性差异，'
         '多模型协同决策比单一模型更可靠。这正是作品名称中'
         '"多模型决策边界对比"的核心教学意图。',
         indent=True)
add_blank_line()

add_body('层次三：AI 智能辅助——大模型驱动的实验解读与问答', bold=True, indent=True)
add_body('本项目以 DeepSeek 大模型为技术基座（同时兼容 Ollama 本地模型'
         '和 SiliconFlow 在线 API），结合 scikit-learn 传统机器学习算法，'
         '实现 AI 分类决策边界可视化与智能解读。', indent=True)
add_blank_line()
add_body('AI 实验解读功能：当学生完成一组实验操作后，系统将当前参数'
         '和四模型的判定结果自动发送至大模型，生成一段通俗易懂的中文'
         '解读，帮助学生理解"为什么会出现这个结果"，实现从"看见现象"'
         '到"理解原理"的认知跃迁。', indent=True)
add_blank_line()
add_body('AI 问答助教功能：学生在实验过程中可以随时向大模型提问，'
         '大模型结合当前实验上下文实时回答。这种"做中学 + 问中学"的'
         '双通道学习模式，让探究过程更加自主和深入。', indent=True)
add_blank_line()
add_body('三个层次层层递进：先「打开黑箱」看见决策边界 → '
         '再「多模型对比」理解算法差异 → '
         '最后「AI 辅助解读」深化认知，形成完整的探究式学习闭环。',
         indent=True)
add_blank_line()

add_heading_para('3. 教育效果', level=3)
add_body('通过本智能体的探究式学习，预期达成以下教育效果：', indent=True)
add_blank_line()
add_bullet('✅ 认知转变：破除"AI 准确率 100% = 绝对安全"的迷思，'
           '建立"AI 也会看走眼"的正确认知')
add_bullet('✅ 概念建构：直观理解"决策边界""特征空间""过拟合"'
           '等核心概念的内含')
add_bullet('✅ 对比思维：通过四模型同台竞技，理解不同算法的'
           '特性差异和各自的优劣势')
add_bullet('✅ 数据意识：通过调节"重叠度""投毒率"，'
           '深刻理解训练数据质量对 AI 模型的根本性影响')
add_bullet('✅ 探究能力：通过"记录—对比—分析—导出"的完整流程，'
           '培养学生的科学探究和数据分析能力')
add_bullet('✅ 批判性思维：通过"共识 vs 分歧"的分析，'
           '培养学生对 AI 决策结果进行审慎判断的意识')
add_blank_line()

# 保存
doc.save(OUTPUT_PATH)
print(f"✅ 操作手册已生成：{OUTPUT_PATH}")
print(f"   主题：{SCENE}")
